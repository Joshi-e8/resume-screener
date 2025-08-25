"""
LLM-based Universal Resume Parser
Uses contextual analysis without hardcoded extraction rules
"""

import json
import os
import time
import hashlib
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional
from pathlib import Path

import aiofiles
import pdfplumber
import PyPDF2
from docx import Document
from loguru import logger

from app.core.config import settings


class LLMResumeParser:
    """Universal resume parser using LLM contextual analysis"""

    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".doc", ".txt"]
        
        # Initialize LLM client based on configuration
        self.llm_client = None
        self._init_llm_client()

    def _init_llm_client(self):
        """Initialize the appropriate LLM client"""
        try:
            provider = getattr(settings, "PROVIDER", "openai").lower()
            
            if provider == "openai":
                from openai import OpenAI
                api_key = getattr(settings, "OPENAI_API_KEY", "")
                base_url = getattr(settings, "OPENAI_BASE_URL", None)
                self.llm_client = OpenAI(api_key=api_key, base_url=base_url)
                self.model = getattr(settings, "OPENAI_MODEL", "gpt-4o")
                self.provider = "openai"
            elif provider == "groq":
                from groq import Groq
                api_key = getattr(settings, "GROQ_API_KEY", "")
                self.llm_client = Groq(api_key=api_key)
                self.model = getattr(settings, "GROQ_MODEL", "llama-3.1-70b-versatile")
                self.provider = "groq"
            else:
                logger.warning(f"Unsupported LLM provider: {provider}")
                self.llm_client = None
                
        except Exception as e:
            logger.error(f"Failed to initialize LLM client: {e}")
            self.llm_client = None

    async def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume using LLM contextual analysis
        """
        start_time = time.time()
        
        # Extract file metadata
        filename = os.path.basename(file_path)
        file_extension = os.path.splitext(filename)[1].lower()
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        
        logger.info(f"[llm_parser] Starting parse for {filename} ({file_size} bytes)")
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Extract raw text
        text_start = time.time()
        raw_text = await self._extract_raw_text(file_path, file_extension)
        text_time = int((time.time() - text_start) * 1000)
        
        if not raw_text or not raw_text.strip():
            logger.warning(f"No text extracted from {filename}")
            return self._create_empty_result(filename, file_extension, file_size)

        logger.info(f"[llm_parser] Text extraction: {text_time}ms, {len(raw_text)} chars")

        # Parse with LLM (check for fast mode setting)
        fast_mode = getattr(settings, "PARSER_LLM_FAST_MODE", True)
        llm_start = time.time()
        parsed_result = await self._parse_with_llm(raw_text, filename, file_extension, file_size, fast_mode)
        llm_time = int((time.time() - llm_start) * 1000)
        
        total_time = int((time.time() - start_time) * 1000)
        logger.info(f"[llm_parser] LLM parsing: {llm_time}ms, Total: {total_time}ms")
        
        # Log parsed data for debugging
        self._log_parsed_data(parsed_result, filename)
        
        return parsed_result

    async def parse_batch_resumes(self, resume_files: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Parse multiple resumes in parallel for better throughput
        """
        import asyncio

        async def parse_single(file_data):
            try:
                if 'file_path' in file_data:
                    return await self.parse_resume(file_data['file_path'])
                else:
                    return await self.parse_resume_from_memory(
                        file_data['content'],
                        file_data['filename'],
                        file_data['extension']
                    )
            except Exception as e:
                logger.error(f"Failed to parse {file_data.get('filename', 'unknown')}: {e}")
                return self._create_empty_result(
                    file_data.get('filename', 'unknown'),
                    file_data.get('extension', '.pdf'),
                    len(file_data.get('content', b''))
                )

        # Process up to 3 resumes in parallel to avoid API rate limits
        semaphore = asyncio.Semaphore(3)

        async def parse_with_semaphore(file_data):
            async with semaphore:
                return await parse_single(file_data)

        tasks = [parse_with_semaphore(file_data) for file_data in resume_files]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Handle any exceptions
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"Batch processing error for file {i}: {result}")
                processed_results.append(self._create_empty_result("error", ".pdf", 0))
            else:
                processed_results.append(result)

        return processed_results

    async def parse_resume_from_memory(self, file_content: bytes, filename: str, file_extension: str) -> Dict[str, Any]:
        """
        Parse resume directly from memory using LLM contextual analysis
        """
        start_time = time.time()
        file_size = len(file_content)
        
        logger.info(f"[llm_parser] Starting memory parse for {filename} ({file_size} bytes)")
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Extract raw text from memory
        text_start = time.time()
        raw_text = await self._extract_raw_text_from_memory(file_content, file_extension)
        text_time = int((time.time() - text_start) * 1000)
        
        if not raw_text or not raw_text.strip():
            logger.warning(f"No text extracted from {filename}")
            return self._create_empty_result(filename, file_extension, file_size)

        logger.info(f"[llm_parser] Memory text extraction: {text_time}ms, {len(raw_text)} chars")

        # Parse with LLM (check for fast mode setting)
        fast_mode = getattr(settings, "PARSER_LLM_FAST_MODE", True)
        llm_start = time.time()
        parsed_result = await self._parse_with_llm(raw_text, filename, file_extension, file_size, fast_mode)
        llm_time = int((time.time() - llm_start) * 1000)
        
        total_time = int((time.time() - start_time) * 1000)
        logger.info(f"[llm_parser] LLM parsing: {llm_time}ms, Total: {total_time}ms")
        
        # Log parsed data for debugging
        self._log_parsed_data(parsed_result, filename)
        
        return parsed_result

    async def _extract_raw_text(self, file_path: str, file_extension: str) -> str:
        """Extract raw text from file"""
        if file_extension == ".pdf":
            return await self._extract_pdf_text(file_path)
        elif file_extension in [".docx", ".doc"]:
            return await self._extract_docx_text(file_path)
        elif file_extension == ".txt":
            return await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    async def _extract_raw_text_from_memory(self, file_content: bytes, file_extension: str) -> str:
        """Extract raw text from file content in memory"""
        if file_extension == ".pdf":
            return await self._extract_pdf_text_from_memory(file_content)
        elif file_extension in [".docx", ".doc"]:
            return await self._extract_docx_text_from_memory(file_content)
        elif file_extension == ".txt":
            return file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _create_empty_result(self, filename: str, file_extension: str, file_size: int) -> Dict[str, Any]:
        """Create empty result when no text can be extracted"""
        return {
            "prompt_passed": False,
            "prompt_metadata": {
                "filename": filename,
                "mime_type": self._get_mime_type(file_extension),
                "file_size_bytes": file_size,
                "source": "llm_parser",
                "ingested_at_iso": datetime.now(timezone.utc).isoformat()
            },
            "document_overview": {
                "detected_language": None,
                "page_count": None,
                "structure_notes": ["No text could be extracted from this file"]
            },
            "sections": {},
            "contact_cluster": {},
            "semantic_highlights": [],
            "quality": {
                "coverage_ratio": 0.0,
                "missing_signals": ["All resume signals missing - no text extracted"],
                "cautions": ["File could not be processed or contains no readable text"]
            },
            "logs": [
                "Step 1: Failed to extract readable text from the provided file.",
                "Step 2: Returning empty result structure."
            ],
            # Legacy compatibility fields
            "raw_text": "",
            "file_type": file_extension,
            "parsed_at": datetime.now(timezone.utc).isoformat(),
            "contact_info": {},
            "skills": [],
            "education": [],
            "experience": [],
            "summary": "",
            "certifications": [],
            "languages": [],
            "projects": [],
            "extraction_warning": "No text could be extracted from this file"
        }

    def _get_mime_type(self, file_extension: str) -> str:
        """Get MIME type from file extension"""
        mime_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain"
        }
        return mime_types.get(file_extension, "application/octet-stream")

    def _log_parsed_data(self, parsed_result: Dict[str, Any], filename: str):
        """Log parsed data for debugging"""
        try:
            # Log key extracted information
            contact_info = parsed_result.get("contact_cluster", {})
            name = contact_info.get("name_text", {}).get("value", "N/A")
            emails = contact_info.get("email_texts", {}).get("values", [])
            phones = contact_info.get("phone_texts", {}).get("values", [])
            
            logger.info(f"[llm_parser] Parsed {filename}: name='{name}', emails={len(emails)}, phones={len(phones)}")
            
            # Log sections found
            sections = parsed_result.get("sections", {})
            section_names = []
            for k, v in sections.items():
                if v:  # If section has content
                    if isinstance(v, dict) and v.get("raw_block"):
                        section_names.append(k)
                    elif isinstance(v, str) and v.strip():  # Handle string sections from fast mode
                        section_names.append(k)
                    elif isinstance(v, list) and v:
                        section_names.append(k)
            logger.info(f"[llm_parser] Sections found: {section_names}")
            
            # Log using structured JSON format
            from app.core.json_logging import log_parsed_resume

            # Extract skills from legacy format
            skills = []
            if "skills" in parsed_result:
                skills = parsed_result["skills"]
            elif "skills_like" in sections:
                skills_text = sections.get("skills_like", {}).get("text", "")
                if skills_text:
                    # Basic skill extraction from text
                    skills = [s.strip() for s in skills_text.split(",") if s.strip()]

            # Extract experience years
            experience_years = parsed_result.get("total_experience_years", 0)

            # Extract education
            education = []
            if "education" in sections:
                edu_text = sections.get("education", {}).get("text", "")
                if edu_text:
                    education = [{"text": edu_text}]

            # Extract quality score
            quality = parsed_result.get("quality", {})
            quality_score = quality.get("coverage_ratio", 0)

            # Extract semantic highlights for additional context
            highlights = parsed_result.get("semantic_highlights", [])
            key_capabilities = [h.get("verbatim", "")[:60] for h in highlights[:3]]

            # Log using structured JSON format
            log_parsed_resume(
                filename=filename,
                candidate_name=name if name != "N/A" else None,
                candidate_email=emails[0] if emails else None,
                candidate_phone=phones[0] if phones else None,
                skills=skills,
                experience_years=experience_years,
                education=education,
                sections_found=section_names,
                quality_score=quality_score,
                key_capabilities=key_capabilities,
                highlights_count=len(highlights),
                emails_count=len(emails),
                phones_count=len(phones)
            )
            
        except Exception as e:
            logger.warning(f"Failed to log parsed data for {filename}: {e}")

    async def _parse_with_llm(self, raw_text: str, filename: str, file_extension: str, file_size: int, fast_mode: bool = True) -> Dict[str, Any]:
        """Parse resume text using 100% NLP approach with intelligent retry mechanisms - NO FALLBACK"""
        if not self.llm_client:
            raise ValueError("LLM client is required for NLP-first parsing. No fallback available.")

        try:
            # Create the parsing prompt - use enhanced NLP approach
            prompt = self._create_enhanced_nlp_prompt(raw_text, filename, file_extension, file_size)

            # Debug: Log the text being sent to AI
            logger.info(f"[llm_parser] Sending {len(raw_text)} chars to AI for {filename}")
            logger.info(f"[llm_parser] Text preview: {raw_text[:200]}...")

            # Call LLM based on provider - optimized for speed
            system_msg = "Extract resume data quickly and accurately. Return only JSON." if fast_mode else "You are an expert resume parser. Analyze the provided resume text using contextual understanding without relying on hardcoded patterns. Return ONLY valid JSON that matches the specified schema."
            max_tokens = 2000 if fast_mode else 4000

            if self.provider == "openai":
                response = self.llm_client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {
                            "role": "system",
                            "content": system_msg
                        },
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ],
                    temperature=0.0,  # Faster with 0 temperature
                    max_tokens=max_tokens,
                    response_format={"type": "json_object"}
                )
                response_text = response.choices[0].message.content
            elif self.provider == "groq":
                response = self.llm_client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {
                            "role": "system",
                            "content": system_msg
                        },
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ],
                    temperature=0.0,  # Faster with 0 temperature
                    max_tokens=max_tokens,
                    response_format={"type": "json_object"}
                )
                response_text = response.choices[0].message.content
            else:
                raise ValueError(f"Unsupported provider: {self.provider}")

            # Log the AI response for debugging (if enabled)
            log_responses = getattr(settings, "LOG_AI_RESPONSES", True)
            if log_responses:
                logger.info("AI response received for resume parsing",
                           event_type="ai_parsing",
                           event="response_received",
                           filename=filename,
                           provider=self.provider,
                           model=self.model,
                           response_length=len(response_text) if response_text else 0,
                           response_preview=response_text[:200] if response_text else None)

            # Debug: Log the raw response
            logger.info(f"[llm_parser] Raw AI response for {filename}: {response_text[:500]}...")

            # Parse JSON response
            parsed_json = json.loads(response_text)

            # Validate response - NO FALLBACK, use intelligent retry instead
            if not parsed_json or not isinstance(parsed_json, dict):
                raise ValueError("AI returned empty or invalid JSON response")

            # For comprehensive NLP approach, we accept any valid JSON structure
            # The AI should be smart enough to provide meaningful data
            if not parsed_json:
                raise ValueError("AI response is completely empty")

            # Add legacy compatibility fields
            legacy_data = self._convert_to_legacy_format(parsed_json, filename)
            parsed_json.update(legacy_data)

            # Add metadata
            parsed_json["raw_text"] = raw_text
            parsed_json["file_type"] = file_extension
            parsed_json["parsed_at"] = datetime.now(timezone.utc).isoformat()
            parsed_json["processing_mode"] = "llm_universal"

            return parsed_json

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM JSON response for {filename}: {e}")
            # Try with a simpler prompt approach
            return await self._retry_with_simpler_prompt(raw_text, filename, file_extension, file_size, e)
        except Exception as e:
            logger.error(f"LLM parsing failed for {filename}: {e}")
            # Try with a simpler prompt approach
            return await self._retry_with_simpler_prompt(raw_text, filename, file_extension, file_size, e)

    async def _retry_with_simpler_prompt(self, raw_text: str, filename: str, file_extension: str, file_size: int, original_error: Exception) -> Dict[str, Any]:
        """Retry with progressively simpler prompts to leverage full NLP capacity"""
        logger.info(f"[llm_parser] Retrying {filename} with simpler prompt due to: {original_error}")

        # Progressive simplification strategies
        strategies = [
            ("basic_structured", self._create_basic_structured_prompt),
            ("minimal_json", self._create_minimal_json_prompt),
            ("ultra_simple", self._create_ultra_simple_prompt)
        ]

        for strategy_name, prompt_creator in strategies:
            try:
                logger.info(f"[llm_parser] Trying {strategy_name} strategy for {filename}")

                # Create simpler prompt
                prompt = prompt_creator(raw_text)

                # Use more conservative settings
                system_msg = "You are a resume parser. Extract key information and return valid JSON only."
                max_tokens = 1500

                # Call LLM with simpler approach
                if self.provider == "openai":
                    response = self.llm_client.chat.completions.create(
                        model=self.model,
                        messages=[
                            {"role": "system", "content": system_msg},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.1,
                        max_tokens=max_tokens,
                        response_format={"type": "json_object"}
                    )
                    response_text = response.choices[0].message.content
                elif self.provider == "groq":
                    response = self.llm_client.chat.completions.create(
                        model=self.model,
                        messages=[
                            {"role": "system", "content": system_msg},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.1,
                        max_tokens=max_tokens,
                        response_format={"type": "json_object"}
                    )
                    response_text = response.choices[0].message.content
                else:
                    raise ValueError(f"Unsupported provider: {self.provider}")

                # Parse and validate
                parsed_json = json.loads(response_text)

                if parsed_json and isinstance(parsed_json, dict):
                    logger.info(f"[llm_parser] SUCCESS with {strategy_name} strategy for {filename}")

                    # Add legacy compatibility
                    legacy_data = self._convert_to_legacy_format(parsed_json, filename)
                    parsed_json.update(legacy_data)

                    # Add metadata
                    parsed_json["raw_text"] = raw_text
                    parsed_json["file_type"] = file_extension
                    parsed_json["parsed_at"] = datetime.now(timezone.utc).isoformat()
                    parsed_json["processing_mode"] = f"nlp_retry_{strategy_name}"

                    return parsed_json

            except Exception as e:
                logger.warning(f"[llm_parser] {strategy_name} strategy failed for {filename}: {e}")
                continue

        # If all strategies fail, raise comprehensive error
        raise ValueError(f"All NLP parsing strategies failed for {filename}. Original error: {original_error}")

    def _create_basic_structured_prompt(self, raw_text: str) -> str:
        """Create a basic structured prompt that focuses on core information"""
        # Clean and limit text
        cleaned_text = self._clean_extracted_text(raw_text)
        limited_text = cleaned_text[:3000]

        return f"""You are an intelligent resume parser focused on COMPREHENSIVE and GRANULAR extraction.

CRITICAL EXTRACTION RULES:
1. BE GRANULAR: "AWS(EC2, S3)" → extract AWS, EC2, S3 as separate skills
2. BE COMPREHENSIVE: Extract ALL technologies, frameworks, tools, and methodologies mentioned
3. INCLUDE SOFT SKILLS: Communication, collaboration, problem-solving, leadership skills
4. FRAMEWORK SPECIFICITY: "Django REST Framework" ≠ "Django" - extract both if mentioned
5. PARENTHETICAL DETAILS: Extract content in parentheses as additional skills
6. IGNORE ARTIFACTS: Handle PDF formatting issues contextually

EXTRACTION SOURCES: Skills sections, job descriptions, project tech stacks, certifications, tools mentioned

Resume text to analyze:
{limited_text}

Return JSON with this structure:
{{
  "contact_info": {{
    "name": "candidate name",
    "email": "email address",
    "phone": "phone number",
    "location": "location"
  }},
  "skills": ["skill1", "skill2", "skill3"],
  "experience": [
    {{
      "title": "job title",
      "company": "company name",
      "duration": "time period",
      "description": "key responsibilities"
    }}
  ],
  "education": [
    {{
      "degree": "degree name",
      "institution": "school name",
      "year": "graduation year"
    }}
  ],
  "professional_summary": "brief summary"
}}"""

    def _create_minimal_json_prompt(self, raw_text: str) -> str:
        """Create minimal prompt for maximum compatibility"""
        limited_text = raw_text[:2000]

        return f"""Extract key info from this resume as JSON:

{limited_text}

JSON format:
{{
  "name": "",
  "email": "",
  "phone": "",
  "skills": [],
  "summary": ""
}}"""

    def _create_ultra_simple_prompt(self, raw_text: str) -> str:
        """Ultra-simple prompt as last resort"""
        limited_text = raw_text[:1500]

        return f"""Resume text: {limited_text}

Return JSON with name, email, phone, and skills only."""

    def _create_universal_prompt(self, raw_text: str, filename: str, file_extension: str, file_size: int) -> str:
        """Create an enhanced NLP-focused resume parsing prompt for maximum accuracy"""
        # Limit text size but keep more content for better accuracy (first 6000 chars)
        limited_text = raw_text[:6000] if len(raw_text) > 6000 else raw_text

        prompt = f"""# Advanced NLP Resume Parser - Extract comprehensive information with high accuracy

## Instructions:
You are an expert resume parser with deep understanding of various resume formats, layouts, and styles.
Use contextual understanding and natural language processing to extract information accurately.

### Key Principles:
1. **Context-Aware Extraction**: Don't rely on section headers - understand content contextually
2. **Format Flexibility**: Handle diverse resume formats (chronological, functional, hybrid, creative)
3. **Semantic Understanding**: Recognize skills, experience, and qualifications regardless of how they're presented
4. **Accuracy Over Speed**: Prioritize correctness and completeness
5. **Evidence-Based**: Always provide evidence spans for extracted information

## Resume Text to Analyze:
```
{limited_text}
```

## File Metadata:
- Filename: {filename}
- File Type: {file_extension}
- File Size: {file_size} bytes
- Processing Time: {datetime.now(timezone.utc).isoformat()}

## Enhanced Output Structure (JSON only):

{{
  "prompt_passed": true,
  "prompt_metadata": {{
    "filename": "{filename}",
    "mime_type": "{self._get_mime_type(file_extension)}",
    "file_size_bytes": {file_size},
    "source": "enhanced_nlp_parser",
    "processing_time": "{datetime.now(timezone.utc).isoformat()}",
    "parser_version": "2.0"
  }},

  "document_analysis": {{
    "detected_language": "<BCP-47 language code (e.g., 'en-US') or null>",
    "estimated_pages": <integer or null>,
    "resume_format": "<chronological|functional|hybrid|creative|academic|other>",
    "content_quality": {{
      "completeness_score": <0.0-1.0>,
      "structure_clarity": <0.0-1.0>,
      "information_density": <0.0-1.0>
    }},
    "layout_notes": [
      "Observations about document structure, formatting, and organization"
    ]
  }},

  "extracted_sections": {{
    "professional_summary": {{
      "content": "<extracted professional summary/objective text>",
      "source_lines": ["<verbatim lines from resume>"],
      "confidence": <0.0-1.0>,
      "section_type": "<summary|objective|profile|about>"
    }},
    "core_skills": {{
      "technical_skills": [
        {{
          "skill": "<skill name>",
          "category": "<programming|framework|tool|database|cloud|etc>",
          "evidence": "<where this skill was mentioned>",
          "confidence": <0.0-1.0>
        }}
      ],
      "soft_skills": [
        {{
          "skill": "<soft skill name>",
          "evidence": "<context where mentioned>",
          "confidence": <0.0-1.0>
        }}
      ],
      "certifications": [
        {{
          "name": "<certification name>",
          "issuer": "<issuing organization>",
          "date": "<date if available>",
          "status": "<active|expired|in_progress|null>"
        }}
      ]
    }},
    "work_experience": [
      {{
        "job_title": "<position title>",
        "company": "<company name>",
        "location": "<work location if mentioned>",
        "duration": {{
          "start_date": "<parsed start date or text>",
          "end_date": "<parsed end date or 'Present'>",
          "duration_text": "<original duration text>"
        }},
        "responsibilities": [
          "<key responsibility or achievement>"
        ],
        "technologies_used": ["<technology>"],
        "achievements": [
          "<quantified achievement or impact>"
        ],
        "confidence": <0.0-1.0>
      }}
    ],
    "education": [
      {{
        "degree": "<degree type and field>",
        "institution": "<school/university name>",
        "location": "<institution location if available>",
        "graduation_date": "<date or year>",
        "gpa": "<GPA if mentioned>",
        "honors": ["<academic honors or distinctions>"],
        "relevant_coursework": ["<relevant courses if listed>"],
        "confidence": <0.0-1.0>
      }}
    ],
    "projects": [
      {{
        "name": "<project name>",
        "description": "<project description>",
        "technologies": ["<technology used>"],
        "duration": "<project timeline>",
        "role": "<your role in project>",
        "outcomes": ["<project results or impact>"],
        "url": "<project URL if available>"
      }}
    ]
  }},

  "contact_information": {{
    "full_name": {{
      "value": "<candidate's full name as it appears>",
      "parsed_components": {{
        "first_name": "<first name>",
        "last_name": "<last name>",
        "middle_name": "<middle name or initial if present>"
      }},
      "evidence_lines": ["<lines where name appears>"],
      "confidence": <0.0-1.0>
    }},
    "email_addresses": [
      {{
        "email": "<email address>",
        "type": "<personal|work|academic|other>",
        "evidence": "<where found in resume>",
        "confidence": <0.0-1.0>
      }}
    ],
    "phone_numbers": [
      {{
        "number": "<phone number as written>",
        "type": "<mobile|home|work|other>",
        "country_code": "<detected country code>",
        "evidence": "<where found>",
        "confidence": <0.0-1.0>
      }}
    ],
    "online_profiles": [
      {{
        "platform": "<LinkedIn|GitHub|Portfolio|Twitter|etc>",
        "url": "<full URL>",
        "username": "<username if extractable>",
        "evidence": "<where found>",
        "confidence": <0.0-1.0>
      }}
    ],
    "location": {{
      "current_location": "<city, state/country as mentioned>",
      "willing_to_relocate": <true|false|null>,
      "remote_work": <true|false|null>,
      "evidence": "<where location info was found>",
      "confidence": <0.0-1.0>
    }}
  }},

  "key_insights": [
    {{
      "category": "<technical_strength|leadership|achievement|domain_expertise|career_progression>",
      "insight": "<specific insight about the candidate>",
      "supporting_evidence": "<verbatim text that supports this insight>",
      "relevance_score": <0.0-1.0>,
      "confidence": <0.0-1.0>
    }}
  ],

  "quality_assessment": {{
    "overall_completeness": <0.0-1.0>,
    "information_richness": <0.0-1.0>,
    "structure_clarity": <0.0-1.0>,
    "missing_elements": [
      "<list of typical resume elements that appear to be missing>"
    ],
    "extraction_challenges": [
      "<any formatting or parsing challenges encountered>"
    ],
    "confidence_factors": [
      "<factors that increase or decrease confidence in the extraction>"
    ]
  }},

  "parsing_metadata": {{
    "extraction_method": "enhanced_nlp",
    "text_length_processed": <number of characters>,
    "sections_identified": <number of distinct sections found>,
    "confidence_distribution": {{
      "high_confidence": <percentage of extractions with >0.8 confidence>,
      "medium_confidence": <percentage with 0.5-0.8 confidence>,
      "low_confidence": <percentage with <0.5 confidence>
    }},
    "processing_notes": [
      "Step 1: Analyzed document structure and identified key sections",
      "Step 2: Extracted contact information using contextual cues",
      "Step 3: Parsed work experience with attention to dates and achievements",
      "Step 4: Identified skills through semantic analysis rather than keyword matching",
      "Step 5: Validated extracted information against evidence in text"
    ]
  }}
}}

Return ONLY the JSON object, no additional text or formatting."""

        return prompt

    def _create_enhanced_nlp_prompt(self, raw_text: str, filename: str, file_extension: str, file_size: int) -> str:
        """Create an enhanced NLP-focused prompt that actually works reliably"""
        # Get configurable text limit
        from app.core.config import settings
        text_limit = getattr(settings, 'PARSER_TEXT_LIMIT', 6000)

        # Clean the text to remove PDF extraction artifacts
        cleaned_text = self._clean_extracted_text(raw_text)

        # Limit text for better processing
        limited_text = cleaned_text[:text_limit] if len(cleaned_text) > text_limit else cleaned_text

        prompt = f"""You are an expert resume parser with advanced contextual understanding.

CRITICAL INSTRUCTIONS:
1. COMPREHENSIVE EXTRACTION: Extract ALL skills, technologies, and tools mentioned - be granular, not general
2. ARTIFACT HANDLING: This text may contain PDF artifacts (Á, µ, à, ¹, ², ³, S, N) - interpret contextually
3. DETAILED PARSING: When you see "AWS(EC2, S3)" extract: AWS, EC2, S3 as separate skills
4. FRAMEWORK SPECIFICITY: "Django REST Framework" is different from "Django" - extract both
5. SOFT SKILLS: Include communication, collaboration, problem-solving skills explicitly mentioned
6. PARENTHETICAL DETAILS: Extract details in parentheses as separate items (e.g., "CI/CD(Github Actions)" → CI/CD, Github Actions)

PARSING STRATEGY:
- Look for skills in: SKILLS sections, Tech Stack lists, job descriptions, project descriptions
- Extract both general (AWS) and specific (EC2, S3) technologies
- Include soft skills, methodologies (Agile), and processes
- Be comprehensive - don't generalize or summarize

Resume Text:
{limited_text}

Extract the following information and return as JSON:

{{
  "prompt_passed": true,
  "contact_info": {{
    "name": "Full name of candidate",
    "email": "Email address",
    "phone": "Phone number",
    "location": "Current location",
    "linkedin": "LinkedIn profile URL if present"
  }},
  "professional_summary": "Professional summary or objective",
  "skills": ["List", "of", "technical", "skills"],
  "experience": [
    {{
      "title": "Job title",
      "company": "Company name",
      "duration": "Employment duration",
      "description": "Key responsibilities and achievements",
      "technologies": ["Technologies", "used"]
    }}
  ],
  "education": [
    {{
      "degree": "Degree type and field",
      "institution": "School/University name",
      "year": "Graduation year",
      "details": "Additional details like GPA, honors"
    }}
  ],
  "projects": [
    {{
      "name": "Project name",
      "description": "Project description",
      "technologies": ["Technologies", "used"]
    }}
  ],
  "certifications": ["List of certifications"],
  "languages": ["List of languages"],
  "key_achievements": ["Notable achievements with quantified results"]
}}

IMPORTANT:
- Return ONLY the JSON object, no explanations
- Use actual data from the resume text
- If information is not available, use empty string or empty array
- Ensure all JSON is properly formatted and valid"""

        return prompt

    def _clean_extracted_text(self, raw_text: str) -> str:
        """Clean text using NLP-aware approach - let AI handle artifacts contextually"""
        import re

        # Minimal cleaning - only remove obvious noise, let NLP handle the rest
        text = raw_text

        # Only remove clearly broken formatting that would confuse any parser
        # Remove excessive whitespace and normalize line breaks
        text = re.sub(r'\s+', ' ', text)  # Multiple spaces -> single space
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple line breaks -> double line break

        # Remove lines that are clearly just formatting noise (very short, non-meaningful)
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            # Only skip completely empty lines or single character artifacts
            if len(line) > 1:  # Keep anything with 2+ characters - let NLP decide if it's meaningful
                cleaned_lines.append(line)

        cleaned_text = '\n'.join(cleaned_lines)

        # Add instruction to AI about handling artifacts
        if self._has_potential_artifacts(cleaned_text):
            # Prepend instruction for AI to handle artifacts contextually
            artifact_instruction = """
[INSTRUCTION: This text may contain PDF extraction artifacts like special characters (Á, µ, à, ¹, ², ³),
formatting symbols (S, N), or broken spacing. Please interpret the content contextually and extract
meaningful information while ignoring obvious formatting artifacts.]

"""
            cleaned_text = artifact_instruction + cleaned_text

        return cleaned_text

    def _has_potential_artifacts(self, text: str) -> bool:
        """Detect if text likely contains PDF extraction artifacts"""
        # Check for common PDF extraction artifacts without hardcoding specific fixes
        artifact_patterns = [
            r'[ÁµàáéíóúÀÈÌÒÙâêîôûäëïöüÿñç]',  # Accented/special chars often from PDF
            r'\b[SN]\s+[A-Z]',  # Single letters followed by words (S Languages, N Tools)
            r'[¹²³⁴⁵⁶⁷⁸⁹⁰]',  # Superscript numbers
            r'[•◦▪▫■□●○]',  # Various bullet point symbols
            r'\b[A-Z]\s*$',  # Single capital letters at end of lines
        ]

        import re
        for pattern in artifact_patterns:
            if re.search(pattern, text):
                return True
        return False

    def _create_fast_prompt(self, raw_text: str, filename: str, file_extension: str, file_size: int) -> str:
        """Create a fast, concise parsing prompt for speed"""
        # Limit text even more for speed (first 3000 chars)
        limited_text = raw_text[:3000] if len(raw_text) > 3000 else raw_text

        prompt = f"""You are a resume parser. Extract information from this resume text and return ONLY valid JSON.

Resume Text:
{limited_text}

Return this exact JSON structure (fill with actual data from the resume):
{{
  "prompt_passed": true,
  "contact_cluster": {{
    "name_text": {{"value": "CANDIDATE_NAME_HERE", "confidence": 0.9}},
    "email_texts": {{"values": ["email@example.com"], "confidence": 0.9}},
    "phone_texts": {{"values": ["+1234567890"], "confidence": 0.9}},
    "location_text": {{"value": "City, State", "confidence": 0.8}}
  }},
  "sections": {{
    "summary": "Professional summary from resume",
    "skills_like": "List of technical skills and technologies",
    "experience_like": "Work experience details",
    "education_like": "Education background"
  }},
  "semantic_highlights": [
    {{"label": "key_skill", "verbatim": "Important skill or technology", "confidence": 0.8}}
  ],
  "quality": {{"coverage_ratio": 0.85}}
}}

IMPORTANT: Return ONLY the JSON object. No explanations, no markdown, no extra text."""

        return prompt

    # FALLBACK REMOVED - 100% NLP APPROACH ONLY
    # All parsing now uses intelligent retry mechanisms with progressive prompt simplification

    def _removed_fallback_method_placeholder(self, raw_text: str, filename: str, file_extension: str, file_size: int) -> Dict[str, Any]:
        """FALLBACK REMOVED - This method should never be called in 100% NLP approach"""
        raise ValueError(f"Fallback method called for {filename}. This should not happen in 100% NLP mode. All parsing should use intelligent retry mechanisms.")

    def _convert_to_legacy_format(self, parsed_json: Dict[str, Any], filename: str = "unknown") -> Dict[str, Any]:
        """Convert enhanced NLP parser output to legacy format for compatibility"""
        legacy_data = {}

        # Handle both new enhanced format and old contact_cluster format
        if "contact_info" in parsed_json:
            # New enhanced format
            contact_info = parsed_json.get("contact_info", {})
            legacy_data["contact_info"] = {
                "name": contact_info.get("name", ""),
                "email": contact_info.get("email", ""),
                "phone": contact_info.get("phone", ""),
                "location": contact_info.get("location", ""),
                "linkedin": contact_info.get("linkedin", "")
            }
        else:
            # Old contact_cluster format
            contact_cluster = parsed_json.get("contact_cluster", {})

            # Extract emails and phones as single values (not lists) for legacy compatibility
            emails = contact_cluster.get("email_texts", {}).get("values", [])
            phones = contact_cluster.get("phone_texts", {}).get("values", [])

            legacy_data["contact_info"] = {
                "name": contact_cluster.get("name_text", {}).get("value", ""),
                "email": emails[0] if emails else "",  # Take first email as string
                "phone": phones[0] if phones else "",  # Take first phone as string
                "location": contact_cluster.get("location_text", {}).get("value", ""),
                "links": contact_cluster.get("link_texts", {}).get("values", [])
            }

        # Handle skills - check for new format first
        if "skills" in parsed_json:
            # New enhanced format - skills are already a list
            legacy_data["skills"] = parsed_json.get("skills", [])
        else:
            # Old format - extract from skills_like section
            sections = parsed_json.get("sections", {})
            skills_section = sections.get("skills_like", {})
            # Handle both dict format {"raw_block": "..."} and string format
            if isinstance(skills_section, dict):
                skills_text = skills_section.get("raw_block", "")
            else:
                skills_text = str(skills_section) if skills_section else ""
            legacy_data["skills"] = self._extract_skills_from_text(skills_text)

        # Debug logging
        logger.info(f"[llm_parser] Legacy conversion for {filename}:")
        logger.info(f"  - Name: {legacy_data['contact_info']['name']}")
        logger.info(f"  - Email: {legacy_data['contact_info']['email']}")
        logger.info(f"  - Skills: {len(legacy_data['skills'])} found")
        logger.info(f"  - Summary length: {len(legacy_data.get('summary', ''))}")

        # Handle summary - check for new format first
        if "professional_summary" in parsed_json:
            # New enhanced format
            legacy_data["summary"] = parsed_json.get("professional_summary", "")
        else:
            # Old format
            sections = parsed_json.get("sections", {})
            summary_section = sections.get("summary", {})
            if isinstance(summary_section, dict):
                legacy_data["summary"] = summary_section.get("raw_block", "")
            else:
                legacy_data["summary"] = str(summary_section) if summary_section else ""

        # Handle experience - check for new format first
        if "experience" in parsed_json:
            # New enhanced format - convert to legacy format
            experience_list = parsed_json.get("experience", [])
            legacy_experience = []
            for exp in experience_list:
                if isinstance(exp, dict):
                    legacy_experience.append({
                        "title": exp.get("title", ""),
                        "company": exp.get("company", ""),
                        "duration": exp.get("duration", ""),
                        "description": exp.get("description", ""),
                        "technologies": exp.get("technologies", [])
                    })
            legacy_data["experience"] = legacy_experience
        else:
            # Old format
            sections = parsed_json.get("sections", {})
            experience_section = sections.get("experience_like", {})
            if isinstance(experience_section, dict):
                experience_text = experience_section.get("raw_block", "")
            else:
                experience_text = str(experience_section) if experience_section else ""
            legacy_data["experience"] = [{"description": experience_text}] if experience_text else []

        # Handle education - check for new format first
        if "education" in parsed_json:
            # New enhanced format
            education_list = parsed_json.get("education", [])
            legacy_education = []
            for edu in education_list:
                if isinstance(edu, dict):
                    legacy_education.append({
                        "degree": edu.get("degree", ""),
                        "institution": edu.get("institution", ""),
                        "year": edu.get("year", ""),
                        "details": edu.get("details", "")
                    })
            legacy_data["education"] = legacy_education
        else:
            # Old format
            sections = parsed_json.get("sections", {})
            education_section = sections.get("education_like", {})
            if isinstance(education_section, dict):
                education_text = education_section.get("raw_block", "")
            else:
                education_text = str(education_section) if education_section else ""
            legacy_data["education"] = [{"description": education_text}] if education_text else []

        # Handle projects - check for new format first
        if "projects" in parsed_json:
            # New enhanced format
            legacy_data["projects"] = parsed_json.get("projects", [])
        else:
            # Old format
            sections = parsed_json.get("sections", {})
            projects_section = sections.get("projects_like", {})
            if isinstance(projects_section, dict):
                projects_text = projects_section.get("raw_block", "")
            else:
                projects_text = str(projects_section) if projects_section else ""
            legacy_data["projects"] = [{"description": projects_text}] if projects_text else []

        # Handle certifications and languages - check for new format first
        if "certifications" in parsed_json:
            legacy_data["certifications"] = parsed_json.get("certifications", [])
        else:
            legacy_data["certifications"] = []

        if "languages" in parsed_json:
            legacy_data["languages"] = parsed_json.get("languages", [])
        else:
            legacy_data["languages"] = []

        # Handle old format other_blocks if new format fields not present
        if not legacy_data["certifications"] and not legacy_data["languages"]:
            sections = parsed_json.get("sections", {})
            other_blocks = sections.get("other_blocks", [])

            for block in other_blocks:
                if isinstance(block, dict):
                    label = block.get("label", "").lower()
                    raw_block = block.get("raw_block", "")

                if "cert" in label or "license" in label:
                    legacy_data["certifications"].append({"name": raw_block})
                elif "lang" in label:
                    legacy_data["languages"].append({"name": raw_block})

        return legacy_data

    def _extract_skills_from_text(self, skills_text: str) -> List[str]:
        """Extract individual skills from skills text block"""
        if not skills_text:
            return []

        # Simple extraction - split by common delimiters
        import re
        skills = []

        # Split by common delimiters
        parts = re.split(r'[,;•\n\r\t]', skills_text)

        for part in parts:
            skill = part.strip()
            if skill and len(skill) > 1 and len(skill) < 50:
                # Remove common prefixes/suffixes
                skill = re.sub(r'^[-•\s]+', '', skill)
                skill = re.sub(r'[-•\s]+$', '', skill)
                if skill:
                    skills.append(skill)

        return skills[:20]  # Limit to 20 skills

    def _extract_skills_dynamically(self, raw_text: str) -> List[str]:
        """Extract skills dynamically without hardcoded lists using pattern recognition"""
        import re
        skills = []

        # Look for skill-like patterns in the text
        lines = raw_text.split('\n')

        for line in lines:
            line = line.strip()

            # Skip empty lines or lines that are too short/long
            if not line or len(line) < 3 or len(line) > 200:
                continue

            # Look for lines that contain skill-like patterns
            # Skills are often in bullet points, comma-separated, or after keywords
            skill_indicators = [
                r'(?:skills?|technologies?|tools?|languages?|frameworks?)[:\s]*(.+)',
                r'(?:proficient|experienced|familiar)\s+(?:in|with)[:\s]*(.+)',
                r'(?:•|-)[\s]*([A-Za-z][A-Za-z0-9\s\+\#\.\-]{1,30})',  # Bullet points
                r'([A-Za-z][A-Za-z0-9\+\#\.\-]{2,20})(?:\s*[,;|]|\s*$)',  # Comma/semicolon separated
            ]

            for pattern in skill_indicators:
                matches = re.finditer(pattern, line, re.IGNORECASE)
                for match in matches:
                    potential_skills = match.group(1) if match.groups() else match.group(0)

                    # Split by common delimiters
                    skill_parts = re.split(r'[,;|•\-\n\r\t]', potential_skills)

                    for part in skill_parts:
                        skill = part.strip()

                        # Filter out non-skill-like text
                        if self._is_likely_skill(skill):
                            skills.append(skill)

        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower not in seen:
                seen.add(skill_lower)
                unique_skills.append(skill)

        # Get configurable skills limit
        from app.core.config import settings
        max_skills = getattr(settings, 'PARSER_MAX_SKILLS', 25)

        return unique_skills[:max_skills]

    def _extract_experience_fallback(self, raw_text: str) -> List[Dict[str, str]]:
        """Extract work experience using pattern recognition (fallback mode)"""
        import re

        experience_entries = []
        lines = raw_text.split('\n')

        # Look for experience section
        in_experience_section = False
        current_job = {}

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect experience section start
            if re.search(r'\b(experience|work|employment|career)\b', line, re.IGNORECASE):
                in_experience_section = True
                continue

            # Stop at next major section
            if in_experience_section and re.search(r'\b(education|projects|skills|certifications)\b', line, re.IGNORECASE):
                if current_job:
                    experience_entries.append(current_job)
                break

            if in_experience_section:
                # Look for job title patterns
                if re.search(r'\b(developer|engineer|manager|analyst|consultant|intern|trainee)\b', line, re.IGNORECASE):
                    if current_job:
                        experience_entries.append(current_job)
                    current_job = {"title": line, "company": "", "duration": "", "description": "", "technologies": []}

                # Look for company names (often in ALL CAPS or followed by location)
                elif re.search(r'^[A-Z\s&]+$', line) and len(line.split()) <= 5:
                    if current_job:
                        current_job["company"] = line

                # Look for dates
                elif re.search(r'\d{4}|\d{2}/\d{4}|present|current', line, re.IGNORECASE):
                    if current_job:
                        current_job["duration"] = line

                # Collect description lines
                elif line.startswith('•') or line.startswith('-') or line.startswith('Á'):
                    if current_job:
                        if current_job["description"]:
                            current_job["description"] += " " + line
                        else:
                            current_job["description"] = line

        # Add last job if exists
        if current_job:
            experience_entries.append(current_job)

        return experience_entries[:5]  # Limit to 5 entries

    def _extract_education_fallback(self, raw_text: str) -> List[Dict[str, str]]:
        """Extract education using pattern recognition (fallback mode)"""
        import re

        education_entries = []
        lines = raw_text.split('\n')

        # Look for education section
        in_education_section = False
        current_edu = {}

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect education section
            if re.search(r'\b(education|academic|qualification|degree)\b', line, re.IGNORECASE):
                in_education_section = True
                continue

            # Stop at next major section
            if in_education_section and re.search(r'\b(experience|projects|skills|certifications)\b', line, re.IGNORECASE):
                if current_edu:
                    education_entries.append(current_edu)
                break

            if in_education_section:
                # Look for degree patterns
                if re.search(r'\b(bachelor|master|phd|diploma|certificate|b\.?tech|m\.?tech|b\.?com|m\.?com|bca|mca)\b', line, re.IGNORECASE):
                    if current_edu:
                        education_entries.append(current_edu)
                    current_edu = {"degree": line, "institution": "", "year": "", "details": ""}

                # Look for institution names (often in title case or ALL CAPS)
                elif re.search(r'\b(college|university|institute|school)\b', line, re.IGNORECASE):
                    if current_edu:
                        current_edu["institution"] = line

                # Look for years
                elif re.search(r'\b(19|20)\d{2}\b', line):
                    if current_edu:
                        current_edu["year"] = line

                # Look for GPA/CGPA
                elif re.search(r'\b(gpa|cgpa|percentage|%)\b', line, re.IGNORECASE):
                    if current_edu:
                        current_edu["details"] = line

        # Add last education if exists
        if current_edu:
            education_entries.append(current_edu)

        return education_entries[:3]  # Limit to 3 entries

    def _extract_summary_fallback(self, raw_text: str) -> str:
        """Extract professional summary using pattern recognition (fallback mode)"""
        import re

        lines = raw_text.split('\n')
        summary_lines = []

        # Look for summary section
        in_summary_section = False

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect summary section
            if re.search(r'\b(summary|objective|profile|about)\b', line, re.IGNORECASE):
                in_summary_section = True
                continue

            # Stop at next major section
            if in_summary_section and re.search(r'\b(experience|education|skills|projects)\b', line, re.IGNORECASE):
                break

            if in_summary_section:
                summary_lines.append(line)

            # If no explicit summary section, take first few descriptive lines
            if not summary_lines and len(line) > 50 and not re.search(r'@|phone|\+\d', line):
                if re.search(r'\b(developer|engineer|experienced|passionate|skilled)\b', line, re.IGNORECASE):
                    summary_lines.append(line)
                    if len(' '.join(summary_lines)) > 200:
                        break

        return ' '.join(summary_lines)[:500]  # Limit to 500 chars

    def _is_likely_skill(self, text: str) -> bool:
        """Determine if text is likely a skill using heuristics (no hardcoded lists)"""
        import re

        if not text or len(text) < 2 or len(text) > 40:
            return False

        # Remove common prefixes/suffixes
        text = re.sub(r'^[-•\s]+', '', text)
        text = re.sub(r'[-•\s]+$', '', text)

        if not text:
            return False

        # Heuristics for skill-like text (no hardcoded skill names)
        skill_patterns = [
            r'^[A-Z][a-z]+(?:\.[a-z]+)*$',  # CamelCase or dotted (e.g., Node.js)
            r'^[A-Z]{2,}$',  # Acronyms (e.g., SQL, AWS, API)
            r'^[A-Za-z]+[\+\#]$',  # Languages with symbols (e.g., C++, C#)
            r'^[A-Za-z]+\s+[A-Za-z]+$',  # Two words (e.g., Machine Learning)
            r'^[A-Za-z]+[-_][A-Za-z]+$',  # Hyphenated/underscored (e.g., React-Native)
        ]

        for pattern in skill_patterns:
            if re.match(pattern, text):
                return True

        # Additional checks for common skill characteristics
        if (text[0].isupper() and  # Starts with capital
            not any(word in text.lower() for word in ['the', 'and', 'or', 'with', 'for', 'in', 'at', 'on']) and  # Not common words
            not re.search(r'\d{4}', text) and  # Not years
            not '@' in text and  # Not email
            not any(char in text for char in '()[]{}') and  # No brackets
            len(text.split()) <= 3):  # Not too many words
            return True

        return False

    # Text extraction methods (reuse from existing parser)
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""

        # Method 1: Try PDFPlumber first (best for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning(f"PDFPlumber failed for {file_path}: {str(e)}")

        # Method 2: Fallback to PyPDF2
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_path}: {str(e)}")

        logger.warning(f"Could not extract text from PDF {file_path}")
        return ""

    async def _extract_pdf_text_from_memory(self, file_content: bytes) -> str:
        """Extract text from PDF file content in memory"""
        try:
            # Try PDFPlumber first
            import io
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                if text_parts:
                    return "\n".join(text_parts)
        except Exception as e:
            logger.warning(f"PDFPlumber failed for memory content: {e}")

        # Fallback to PyPDF2
        try:
            import io
            from PyPDF2 import PdfReader

            pdf_reader = PdfReader(io.BytesIO(file_content))
            text_parts = []

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            if text_parts:
                return "\n".join(text_parts)
        except Exception as e:
            logger.warning(f"PyPDF2 failed for memory content: {e}")

        logger.warning("Could not extract text from PDF memory content")
        return ""

    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n".join(text_parts) if text_parts else ""

        except Exception as e:
            logger.warning(f"DOCX extraction failed for {file_path}: {str(e)}")
            return ""

    async def _extract_docx_text_from_memory(self, file_content: bytes) -> str:
        """Extract text from DOCX file content in memory"""
        try:
            import io
            doc = Document(io.BytesIO(file_content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return "\n".join(text_parts)
        except Exception as e:
            logger.warning(f"DOCX extraction failed for memory content: {e}")
            return ""

    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as file:
                text = await file.read()
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract TXT text: {str(e)}")
            return ""
