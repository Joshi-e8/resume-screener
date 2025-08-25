"""
Resume parsing service using PDFPlumber and other libraries
"""

import os
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

import aiofiles
import pdfplumber
import PyPDF2
from docx import Document



class ResumeParser:
    """Service for parsing resumes from various file formats"""

    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".doc", ".txt"]

        # Common patterns for extracting information
        self.email_pattern = re.compile(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        )
        self.phone_pattern = re.compile(
            r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
        )
        self.linkedin_pattern = re.compile(r"linkedin\.com/in/[\w-]+", re.IGNORECASE)
        self.github_pattern = re.compile(r"github\.com/[\w-]+", re.IGNORECASE)

        # Skills patterns (can be expanded)
        self.tech_skills = [
            "python",
            "java",
            "javascript",
            "typescript",
            "react",
            "angular",
            "vue",
            "node.js",
            "express",
            "django",
            "flask",
            "fastapi",
            "spring",
            "hibernate",
            "sql",
            "mysql",
            "postgresql",
            "mongodb",
            "redis",
            "elasticsearch",
            "aws",
            "azure",
            "gcp",
            "docker",
            "kubernetes",
            "jenkins",
            "git",
            "html",
            "css",
            "sass",
            "bootstrap",
            "tailwind",
            "figma",
            "photoshop",
            "machine learning",
            "ai",
            "data science",
            "pandas",
            "numpy",
            "tensorflow",
            "pytorch",
            "scikit-learn",
            "r",
            "matlab",
            "tableau",
            "power bi",
        ]

        # Education keywords
        self.education_keywords = [
            "bachelor",
            "master",
            "phd",
            "doctorate",
            "degree",
            "university",
            "college",
            "institute",
            "school",
            "education",
            "graduated",
            "gpa",
        ]

        # Experience keywords
        self.experience_keywords = [
            "experience",
            "work",
            "employment",
            "job",
            "position",
            "role",
            "responsibilities",
            "achievements",
            "projects",
            "internship",
        ]

        # Pre-compile regex patterns for better performance
        self.job_title_pattern = re.compile(
            r"(manager|engineer|developer|analyst|specialist|coordinator|director|lead|senior|junior)",
            re.IGNORECASE
        )
        self.year_pattern = re.compile(r"(19|20)\d{2}")
        self.skills_section_pattern = re.compile(
            r"(?:skills?|technologies?|technical skills?)[:\s]*(.*?)(?:\n\n|\n[A-Z]|$)",
            re.IGNORECASE | re.DOTALL
        )

    async def parse_resume_from_memory(self, file_content: bytes, filename: str, file_extension: str) -> Dict[str, Any]:
        """
        Parse resume directly from memory for high-performance bulk processing
        """
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Try universal LLM parser first if enabled (highest priority)
        use_universal_llm = False
        try:
            from app.core.config import settings
            setting_value = getattr(settings, "PARSER_USE_UNIVERSAL_LLM", False)
            if isinstance(setting_value, bool):
                use_universal_llm = setting_value
            else:
                use_universal_llm = bool(int(str(setting_value or "0")))
        except Exception:
            use_universal_llm = False

        if use_universal_llm:
            try:
                from app.services.llm_resume_parser import LLMResumeParser
                llm_parser = LLMResumeParser()
                result = await llm_parser.parse_resume_from_memory(file_content, filename, file_extension)
                print(f"✅ Universal LLM parser completed for {filename} (memory)")
                return result
            except Exception as e:
                print(f"❌ Universal LLM parser failed for {filename}, falling back to fast mode: {e}")
                import traceback
                traceback.print_exc()

        # Extract text based on file type directly from memory
        if file_extension == ".pdf":
            text = await self._extract_pdf_text_from_memory(file_content)
        elif file_extension in [".docx", ".doc"]:
            text = await self._extract_docx_text_from_memory(file_content)
        elif file_extension == ".txt":
            text = file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Handle empty text gracefully
        if not text or not text.strip():
            return {
                "raw_text": "",
                "file_type": file_extension,
                "parsed_at": datetime.now(timezone.utc).isoformat(),
                "contact_info": {},
                "skills": [],
                "education": [],
                "experience": [],
                "summary": "",
                "certifications": [],
                "languages": [],
                "projects": [],
                "extraction_warning": "No text could be extracted from this file"
            }

        # Limit text for fast processing
        if len(text) > 5000:
            text = text[:5000]

        # Parse essential data for maximum speed
        parsed_data = await self._parse_text_content_ultra_fast(text)
        parsed_data["raw_text"] = text
        parsed_data["file_type"] = file_extension
        parsed_data["parsed_at"] = datetime.now(timezone.utc).isoformat()
        parsed_data["processing_mode"] = "fast_bulk"

        return parsed_data

    async def _parse_text_content_ultra_fast(self, text: str) -> Dict[str, Any]:
        """
        Ultra-fast parsing - only extract essential information for maximum speed
        """
        # Skip expensive operations, only extract basics
        lines = text.split('\n')[:50]  # Only process first 50 lines for speed
        cleaned_text = ' '.join(line.strip() for line in lines if line.strip())

        # Only extract the most essential information
        parsed_data = {
            "contact_info": self._extract_contact_info_fast(cleaned_text),
            "skills": self._extract_skills_fast(cleaned_text),
            "education": [],  # Skip for speed
            "experience": [],  # Skip for speed
            "summary": "",  # Skip for speed
            "certifications": [],  # Skip for speed
            "languages": [],  # Skip for speed
            "projects": []  # Skip for speed
        }

        return parsed_data

    def _extract_contact_info_fast(self, text: str) -> Dict[str, Any]:
        """
        Fast contact info extraction - only email and phone
        """
        import re

        contact_info = {}

        # Extract email (simple regex for speed)
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact_info["email"] = email_match.group()

        # Extract phone (simple regex for speed)
        phone_match = re.search(r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b', text)
        if phone_match:
            contact_info["phone"] = phone_match.group()

        return contact_info

    def _extract_skills_fast(self, text: str) -> List[str]:
        """
        Fast skills extraction - only common tech skills
        """
        # Common tech skills for fast matching
        common_skills = [
            'Python', 'Java', 'JavaScript', 'React', 'Node.js', 'SQL', 'AWS', 'Docker',
            'Git', 'HTML', 'CSS', 'TypeScript', 'MongoDB', 'PostgreSQL', 'Redis',
            'Kubernetes', 'Linux', 'REST', 'API', 'Machine Learning', 'AI'
        ]

        text_upper = text.upper()
        found_skills = []

        for skill in common_skills:
            if skill.upper() in text_upper:
                found_skills.append(skill)

        return found_skills[:10]  # Limit to 10 skills for speed

    async def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume from file path and extract structured data
        """
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Try universal LLM parser first if enabled (highest priority)
        use_universal_llm = False
        try:
            from app.core.config import settings
            setting_value = getattr(settings, "PARSER_USE_UNIVERSAL_LLM", False)
            if isinstance(setting_value, bool):
                use_universal_llm = setting_value
            else:
                use_universal_llm = bool(int(str(setting_value or "0")))
        except Exception:
            use_universal_llm = False

        if use_universal_llm:
            try:
                from app.services.llm_resume_parser import LLMResumeParser
                llm_parser = LLMResumeParser()
                result = await llm_parser.parse_resume(file_path)
                print(f"✅ Universal LLM parser completed for {file_path}")
                return result
            except Exception as e:
                print(f"❌ Universal LLM parser failed, falling back to orchestrator: {e}")
                import traceback
                traceback.print_exc()

        # Use NLP-first approach instead of rule-based orchestrator
        use_nlp_first = True
        try:
            from app.core.config import settings
            setting_value = getattr(settings, "PARSER_USE_NLP_FIRST", True)
            if isinstance(setting_value, bool):
                use_nlp_first = setting_value
            else:
                use_nlp_first = bool(int(str(setting_value or "1")))
        except Exception:
            use_nlp_first = True

        if use_nlp_first:
            try:
                from app.services.llm_resume_parser import LLMResumeParser
                import os as _os
                llm_parser = LLMResumeParser()

                # Get file size for metadata
                size = 0
                try:
                    size = _os.path.getsize(file_path)
                except Exception:
                    pass

                # Use LLM parser for accurate extraction
                result = await llm_parser.parse_resume(file_path)

                # Ensure compatibility with existing format
                if result and isinstance(result, dict):
                    result["parsed_at"] = datetime.now(timezone.utc).isoformat()
                    result["processing_mode"] = "nlp_first"
                    return result
                else:
                    print(f"Warning: NLP parser returned invalid result for {file_path}, falling back to legacy parser")

            except Exception as e:
                print(f"NLP parser failed, falling back to legacy: {e}")

        # Legacy extraction path
        if file_extension == ".pdf":
            text = await self._extract_pdf_text(file_path)
        elif file_extension in [".docx", ".doc"]:
            text = await self._extract_docx_text(file_path)
        elif file_extension == ".txt":
            text = await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Handle empty text gracefully
        if not text or not text.strip():
            print(f"Warning: No text extracted from {file_path}")
            return {
                "raw_text": "",
                "file_type": file_extension,
                "parsed_at": datetime.now(timezone.utc).isoformat(),
                "contact_info": {},
                "skills": [],
                "education": [],
                "experience": [],
                "summary": "",
                "certifications": [],
                "languages": [],
                "projects": [],
                "extraction_warning": "No text could be extracted from this file"
            }

        # Limit text size to prevent performance issues (max 20KB for fast processing)
        fast_mode = False
        if len(text) > 20000:
            print(f"Warning: Large text file ({len(text)} chars), enabling fast mode for {file_path}")
            text = text[:20000]
            fast_mode = True

        # Parse structured data from text
        parsed_data = await self._parse_text_content(text, fast_mode)
        parsed_data["raw_text"] = text
        parsed_data["file_type"] = file_extension
        parsed_data["parsed_at"] = datetime.now(timezone.utc).isoformat()
        if fast_mode:
            parsed_data["processing_mode"] = "fast"

        return parsed_data

    async def _extract_pdf_text_from_memory(self, file_content: bytes) -> str:
        """
        Extract text from PDF file content in memory (faster)
        """
        try:
            # Try PDFPlumber first (most reliable)
            import io
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                if text_parts:
                    return "\n".join(text_parts)
        except Exception as e:
            print(f"PDFPlumber failed for memory content: {e}")

        # Fallback to PyPDF2
        try:
            import io
            from PyPDF2 import PdfReader

            pdf_reader = PdfReader(io.BytesIO(file_content))
            text_parts = []

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            if text_parts:
                return "\n".join(text_parts)
        except Exception as e:
            print(f"PyPDF2 failed for memory content: {e}")

        print("Warning: Could not extract text from PDF memory content, returning empty string")
        return ""

    async def _extract_docx_text_from_memory(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file content in memory (faster)
        """
        try:
            import io
            from docx import Document

            doc = Document(io.BytesIO(file_content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return "\n".join(text_parts)
        except Exception as e:
            print(f"DOCX extraction failed for memory content: {e}")
            return ""

    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""

        # Method 1: Try PDFPlumber first (best for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text.strip()
        except Exception as e:
            print(f"PDFPlumber failed for {file_path}: {str(e)}")

        # Method 2: Fallback to PyPDF2
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text.strip()
        except Exception as e:
            print(f"PyPDF2 failed for {file_path}: {str(e)}")

        # Method 3: Try with different PyPDF2 settings
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file, strict=False)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text.strip()
        except Exception as e:
            print(f"PyPDF2 (non-strict) failed for {file_path}: {str(e)}")

        # If all methods fail, return empty string instead of raising exception
        print(f"Warning: Could not extract text from PDF {file_path}, returning empty string")
        return ""

    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file with improved error handling"""
        text = ""

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

            if text.strip():
                return text.strip()
            else:
                print(f"Warning: No text content found in DOCX {file_path}")
                return ""

        except Exception as e:
            print(f"DOCX extraction failed for {file_path}: {str(e)}")
            # Return empty string instead of raising exception
            return ""

    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as file:
                text = await file.read()
            return text.strip()
        except Exception as e:  # noqa: E722
            raise Exception(f"Failed to extract TXT text: {str(e)}")

    async def _parse_text_content(self, text: str, fast_mode: bool = False) -> Dict[str, Any]:
        """Parse structured data from resume text"""
        import time

        # Clean and normalize text
        start_time = time.time()
        cleaned_text = self._clean_text(text)
        lines = cleaned_text.split("\n")

        # Limit number of lines to process for performance (max 500 lines)
        if len(lines) > 500:
            print(f"Warning: Large file with {len(lines)} lines, limiting to 500 lines for processing")
            lines = lines[:500]

        print(f"Text cleaning: {int((time.time() - start_time) * 1000)}ms, {len(lines)} lines")

        parsed_data = {}

        # Time each extraction method
        start_time = time.time()
        parsed_data["contact_info"] = self._extract_contact_info(cleaned_text)
        print(f"Contact info extraction: {int((time.time() - start_time) * 1000)}ms")

        start_time = time.time()
        skills = self._extract_skills_nlp(cleaned_text) or self._extract_skills(cleaned_text)
        parsed_data["skills"] = skills
        print(f"Skills extraction: {int((time.time() - start_time) * 1000)}ms")

        start_time = time.time()
        parsed_data["education"] = self._extract_education(lines)
        print(f"Education extraction: {int((time.time() - start_time) * 1000)}ms")

        if fast_mode:
            # In fast mode, skip expensive operations
            print("Fast mode: skipping detailed experience, summary, certifications, languages, and projects extraction")
            parsed_data["experience"] = []
            parsed_data["summary"] = ""
            parsed_data["certifications"] = []
            parsed_data["languages"] = []
            parsed_data["projects"] = []
        else:
            start_time = time.time()
            parsed_data["experience"] = self._extract_experience(lines)
            print(f"Experience extraction: {int((time.time() - start_time) * 1000)}ms")

            # Try to extract a probable title from the first few lines
            try:
                head = " ".join(lines[:5])[:200].lower()
                titles = [
                    "python developer", "backend developer", "software engineer",
                    "backend engineer", "full stack developer", "full-stack developer"
                ]
                for t in titles:
                    if t in head:
                        parsed_data["title"] = t.title()
                        break
            except Exception:
                pass

            start_time = time.time()
            parsed_data["summary"] = self._extract_summary(lines)
            print(f"Summary extraction: {int((time.time() - start_time) * 1000)}ms")

            start_time = time.time()
            parsed_data["certifications"] = self._extract_certifications(cleaned_text)
            print(f"Certifications extraction: {int((time.time() - start_time) * 1000)}ms")

            start_time = time.time()
            parsed_data["languages"] = self._extract_languages(cleaned_text)
            print(f"Languages extraction: {int((time.time() - start_time) * 1000)}ms")

            start_time = time.time()
            parsed_data["projects"] = self._extract_projects(lines)
            print(f"Projects extraction: {int((time.time() - start_time) * 1000)}ms")

        return parsed_data

    def _map_orchestrator_to_legacy(self, obj: Dict[str, Any], file_extension: str) -> Dict[str, Any]:
        """Map orchestrator output to legacy parsed schema expected by current consumers."""
        # Contact info
        cand = obj.get("candidate") or {}
        contact = {
            "name": cand.get("name"),
            "email": (cand.get("emails") or [None])[0],
            "phone": (cand.get("phones") or [None])[0],
            "linkedin": (cand.get("links") or {}).get("linkedin"),
            "github": (cand.get("links") or {}).get("github"),
            "location": (cand.get("location") or {}).get("raw") or None,
        }
        # Skills (legacy expects simple list)
        skills_struct = obj.get("skills") or []
        skills = [s.get("name") for s in skills_struct if isinstance(s, dict) and s.get("name")]
        # Education simple list of objects
        education = obj.get("education") or []
        # Experience minimal mapping to legacy list
        experience = obj.get("experience") or []
        # Summary
        summary = obj.get("summary") or ""
        # Languages/certs/projects
        languages = [l.get("name") for l in (obj.get("languages") or []) if isinstance(l, dict) and l.get("name")]
        certs = obj.get("certifications") or []
        projects = obj.get("projects") or []
        return {
            "raw_text": "",
            "file_type": file_extension,
            "contact_info": contact,
            "skills": skills,
            "education": education,
            "experience": experience,
            "summary": summary,
            "certifications": certs,
            "languages": languages,
            "projects": projects,
        }


    def _clean_text(self, text: str) -> str:
        """Clean and normalize text while preserving newlines for section parsing."""
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse spaces and tabs but keep newlines
        text = re.sub(r"[\t\x0b\x0c\f]+", " ", text)
        text = re.sub(r"[ ]{2,}", " ", text)
        # Collapse excessive newlines to single
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Keep punctuation useful for skills (.,-/+ and commas for splitting)
        text = re.sub(r"[^\w\s@\.,\-/\+]", " ", text)
        return text.strip()

    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information"""
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "location": None,
        }

        # Extract email
        email_match = self.email_pattern.search(text)
        if email_match:
            contact_info["email"] = email_match.group()

        # Extract phone
        phone_match = self.phone_pattern.search(text)
        if phone_match:
            contact_info["phone"] = phone_match.group()

        # Extract LinkedIn
        linkedin_match = self.linkedin_pattern.search(text)
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group()

        # Extract GitHub
        github_match = self.github_pattern.search(text)
        if github_match:
            contact_info["github"] = github_match.group()


    def _extract_skills_nlp(self, text: str) -> List[str]:
        """
        NLP-based skills/keyword extraction using spaCy if available.
        Returns a cleaned, deduplicated list of plausible skills/technologies/keywords.
        """
        try:
            import re
            nlp = None
            try:
                import spacy
                try:
                    nlp = spacy.load("en_core_web_sm")
                except Exception:
                    # Fall back to a blank English pipeline if model is missing
                    nlp = spacy.blank("en")
                    if "sentencizer" not in nlp.pipe_names:
                        nlp.add_pipe("sentencizer")
            except Exception:
                nlp = None

            text = text or ""
            if not text.strip():
                return []

            # Limit for performance
            sample = text[:15000]
            doc = nlp(sample) if nlp is not None else None

            candidates: List[str] = []

            # 1) Noun chunks as candidate skills/phrases (when model available)
            if doc is not None and hasattr(doc, "noun_chunks"):
                for chunk in doc.noun_chunks:
                    tok = chunk.text.strip()
                    if len(tok) >= 3 and len(tok.split()) <= 5:
                        candidates.append(tok)

            # 2) Named entities (ORG/PRODUCT often capture tech/product names)
            if doc is not None:
                for ent in getattr(doc, "ents", []):
                    if ent.label_ in ("ORG", "PRODUCT", "WORK_OF_ART"):
                        tok = ent.text.strip()
                        if len(tok) >= 3:
                            candidates.append(tok)

            # 3) Token pattern capturing dotted/slashed tech tokens (Next.js, CI/CD)
            for m in re.finditer(r"\b[A-Za-z][A-Za-z0-9\.\+\/#-]{2,}\b", sample):
                tok = m.group(0)
                if any(c in tok for c in [".", "/", "+", "-"]):
                    candidates.append(tok)

            # Clean and dedupe
            cleaned: List[str] = []
            seen = set()
            for s in candidates:
                s = re.sub(r"\s+", " ", s).strip()
                if len(s) <= 2:
                    continue
                # Skip long multi-word phrases unless they include tech punctuation
                if len(s.split()) > 5 and not any(c in s for c in [".", "/", "+", "-"]):
                    continue
                # Filter obvious fragments unless allowlisted
                toks = s.split()
                if any(len(t) <= 2 for t in toks):
                    if not any(x in s for x in ["CI/CD", "C++", "C#", ".js", "S3", "EC2"]):
                        if not (s.isupper() and len(s) <= 5):
                            continue
                key = s.lower()
                if key not in seen:
                    seen.add(key)
                    cleaned.append(s)

            # Prioritize tokens appearing in a skills section
            m = self.skills_section_pattern.search(text)
            if m:
                sec = (m.group(1) or "").lower()
                cleaned.sort(key=lambda t: 0 if t.lower() in sec else 1)

            return cleaned[:30]
        except Exception:
            return []

    def _extract_skills(self, text: str) -> List[str]:
        """Domain-agnostic skills/keywords extraction with noise suppression and proper tokenization."""
        found: List[str] = []
        import re

        def add_skill(s: str):
            s_norm = re.sub(r"\s+", " ", s).strip()
            if not s_norm:
                return
            if len(s_norm) <= 2:
                return
            # Avoid obvious fragments unless common allowlist
            toks = s_norm.split()
            if any(len(t) <= 2 for t in toks):
                if not any(x in s_norm for x in ["CI/CD", "C++", "C#", ".js", "S3", "EC2"]):
                    if not (s_norm.isupper() and len(s_norm) <= 5):
                        return
            if s_norm not in found:
                found.append(s_norm)

        # 1) Prefer items from an explicit Skills section
        m = self.skills_section_pattern.search(text)
        if m:
            skills_text = m.group(1)
            parts = re.split(r"[,;\|\n\r\t•\-]+", skills_text)
            for p in parts:
                token = p.strip()
                if not token:
                    continue
                # keep short phrases (<=5 words) and tokens with tech punctuation
                if len(token.split()) <= 5 or any(c in token for c in ["/", ".", "+", "-"]):
                    add_skill(token)

        # 2) Also capture globally any token with tech/business punctuation
        for m in re.finditer(r"\b[A-Za-z][A-Za-z0-9\.\+\/#-]{2,}\b", text):
            tok = m.group(0)
            if any(c in tok for c in [".", "/", "+", "-"]):
                add_skill(tok)

        return found[:30]

    def _extract_education(self, lines: List[str]) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        education_section = False
        current_education = {}

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if we're in education section
            if any(keyword in line.lower() for keyword in self.education_keywords):
                education_section = True

                # Look for degree patterns
                degree_pattern = r"(bachelor|master|phd|doctorate|b\.?s\.?|m\.?s\.?|m\.?a\.?|ph\.?d\.?)"
                degree_match = re.search(degree_pattern, line, re.IGNORECASE)

                if degree_match:
                    if current_education:
                        education.append(current_education)

                    current_education = {
                        "degree": line,
                        "institution": "",
                        "year": "",
                        "field": "",
                    }

            elif education_section and line:
                # Look for years
                year_pattern = r"(19|20)\d{2}"
                year_match = re.search(year_pattern, line)
                if year_match and "year" not in current_education:
                    current_education["year"] = year_match.group()

                # If it looks like an institution name
                if not current_education.get("institution") and len(line) > 10:
                    current_education["institution"] = line

        if current_education:
            education.append(current_education)

        return education

    def _extract_experience(self, lines: List[str]) -> List[Dict[str, str]]:
        """Extract work experience with performance optimization"""
        experience = []
        experience_section = False
        current_job = {}

        # Limit processing to first 200 lines for performance
        lines_to_process = lines[:200] if len(lines) > 200 else lines

        for line in lines_to_process:
            line = line.strip()
            if not line:
                continue

            # Check if we're in experience section
            if any(keyword in line.lower() for keyword in self.experience_keywords):
                experience_section = True

            elif experience_section:
                # Look for job titles and companies using pre-compiled pattern
                if self.job_title_pattern.search(line):
                    if current_job:
                        experience.append(current_job)

                    current_job = {
                        "title": line,
                        "company": "",
                        "duration": "",
                        "description": "",
                    }

                # Look for years/duration using pre-compiled pattern
                elif self.year_pattern.search(line) and current_job:
                    current_job["duration"] = line

                # Add to description (limit description length)
                elif current_job and len(line) > 20:
                    if len(current_job.get("description", "")) < 500:  # Limit description length
                        if current_job["description"]:
                            current_job["description"] += " " + line
                        else:
                            current_job["description"] = line

        if current_job:
            experience.append(current_job)

        return experience

    def _extract_summary(self, lines: List[str]) -> str:
        """Extract professional summary"""
        summary_keywords = ["summary", "objective", "profile", "about"]

        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in summary_keywords):
                # Get next few lines as summary
                summary_lines = []
                for j in range(i + 1, min(i + 5, len(lines))):
                    if lines[j].strip() and not any(
                        keyword in lines[j].lower()
                        for keyword in self.experience_keywords
                        + self.education_keywords
                    ):
                        summary_lines.append(lines[j].strip())
                    else:
                        break

                return " ".join(summary_lines)

        return ""

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        cert_keywords = ["certification", "certified", "certificate"]
        certifications = []

        for keyword in cert_keywords:
            pattern = rf"{keyword}[:\s]*(.*?)(?:\n|$)"
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if match.strip():
                    certifications.append(match.strip())

        return certifications

    def _extract_languages(self, text: str) -> List[str]:
        """Extract languages"""
        language_pattern = r"(?:languages?)[:\s]*(.*?)(?:\n\n|\n[A-Z]|$)"
        language_match = re.search(language_pattern, text, re.IGNORECASE | re.DOTALL)

        if language_match:
            languages_text = language_match.group(1)
            languages = re.findall(r"([A-Z][a-z]+)", languages_text)
            return [lang for lang in languages if len(lang) > 2]

        return []

    def _extract_projects(self, lines: List[str]) -> List[Dict[str, str]]:
        """Extract project information"""
        projects = []
        project_keywords = ["project", "portfolio"]
        project_section = False
        current_project = {}

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if any(keyword in line.lower() for keyword in project_keywords):
                project_section = True

            elif project_section and line:
                if current_project and (line.startswith("•") or line.startswith("-")):
                    if current_project["description"]:
                        current_project["description"] += " " + line
                    else:
                        current_project["description"] = line
                else:
                    if current_project:
                        projects.append(current_project)

                    current_project = {
                        "name": line,
                        "description": "",
                        "technologies": "",
                    }

        if current_project:
            projects.append(current_project)

        return projects
